"""The document writer can read the rubric template at ./rubric.docx and build
the output document of all student results, which will be returned to students."""

from copy import deepcopy
import logging
from typing import Sequence

from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from .utils import BASE_DIR
from .entities import Page


logger = logging.getLogger(__name__)


TEMPLATE = BASE_DIR / "grader" / "assets" / "rubric.docx"
FIFTH_TEMPLATE = BASE_DIR / "grader" / "assets" / "rubric_fifth.docx"


class DocWriter:
    # for the purposes of shading cells, the grade on the left corresponds to
    # shading the column on the right
    grade_to_col_mapping = {0: 1, 10: 2, 15: 3, 20: 4}

    def __init__(self, type_=None):
        self.type = type_

        if self.type != "fifth":
            self.doc = Document(TEMPLATE)
        else:
            self.doc = Document(FIFTH_TEMPLATE)

    def add_page(self, page: Page):
        self.doc.add_page_break()
        self.doc.add_heading("Virtual Piano Project Feedback")
        par = self.doc.add_paragraph(f"Name: {page.student.name}")
        par._p.addnext(self._template_table)

        # fifth graders don't have this in their template
        if self.type != "fifth":
            self._insert_grade("<week 19>", page.wk_19_grade)

        self._insert_grade("<week 20>", page.wk_20_grade)
        self._insert_grade("<week 21>", page.wk_21_grade)
        self._replace_table_text(
            "<total>", str(page.total_grade), search_tables=[self.doc.tables[-1]]
        )

        if page.notes is not None:
            self.doc.add_paragraph(f"Notes: {page.notes}")

    def add_pages(self, pages: Sequence[Page]):
        for page in pages:
            self.add_page(page)
            logger.debug("wrote page for %s", page.student.name)

    def _insert_grade(self, template_target: str, value: int):
        _, i_row, _ = self._replace_table_text(
            template_target, str(value), search_tables=[self.doc.tables[-1]]
        )
        i_cell = self.grade_to_col_mapping[value]
        cell = self.doc.tables[-1].rows[i_row].cells[i_cell]
        self.highlight_cell(cell)

    def _replace_table_text(self, search, replace: str, search_tables=None):
        """Returns the index of the table, row, and cell that the text was
        found in. If `search_tables` is specified, the returned table index
        refers to the passed-in list, not tables in the whole document."""
        if search_tables is None:
            search_tables = self.doc.tables
        for i, table in enumerate(search_tables):
            for j, row in enumerate(table.rows):
                for k, paragraph in enumerate(row.cells):
                    if search in paragraph.text:
                        # cursed, but this seems to replace the paragraph's
                        # inner text, preventing the formatting from being
                        # changed, and maintaining the text position formatting
                        # as being in the middle of the cell
                        paragraph.paragraphs[0].text = replace
                        return i, j, k

        raise Exception(f"could not perform replacement for {search}")

    @property
    def _template_table(self):
        return deepcopy(self.doc.tables[0]._tbl)

    @staticmethod
    def highlight_cell(cell, color="CCCCCC"):
        shading_elm_1 = parse_xml(
            r'<w:shd {nsdecl} w:fill="{color}"/>'.format(
                nsdecl=nsdecls("w"), color=color
            )
        )
        cell._tc.get_or_add_tcPr().append(shading_elm_1)
